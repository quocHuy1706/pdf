from pathlib import Path
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import html
import os


def _safe_name(name: str) -> str:
    return "".join(c if c.isalnum() or c in "-_" else "_" for c in name)[:80]


def export_exam_docx(exam, export_dir: str = "exports") -> str:
    Path(export_dir).mkdir(exist_ok=True)
    filename = f"de_thi_{exam.id}_{_safe_name(exam.title)}.docx"
    path = Path(export_dir) / filename

    doc = DocxDocument()
    doc.add_heading(exam.title, level=1)
    doc.add_paragraph(f"Mức độ: {exam.difficulty}")
    doc.add_paragraph("Họ và tên: ........................................")
    doc.add_paragraph("Lớp: ...............................................")

    for idx, q in enumerate(exam.questions, start=1):
        doc.add_paragraph(f"Câu {idx}. {q.content}")
        doc.add_paragraph(f"A. {q.option_a}")
        doc.add_paragraph(f"B. {q.option_b}")
        doc.add_paragraph(f"C. {q.option_c}")
        doc.add_paragraph(f"D. {q.option_d}")
        doc.add_paragraph("")

    doc.add_page_break()
    doc.add_heading("Đáp án", level=2)
    answers = "; ".join([f"Câu {i}: {q.correct_answer}" for i, q in enumerate(exam.questions, start=1)])
    doc.add_paragraph(answers)
    doc.save(path)
    return str(path)


def export_exam_pdf(exam, export_dir: str = "exports") -> str:
    Path(export_dir).mkdir(exist_ok=True)
    filename = f"de_thi_{exam.id}_{_safe_name(exam.title)}.pdf"
    path = Path(export_dir) / filename

    # ReportLab mặc định hỗ trợ tiếng Anh tốt. Nếu máy có font DejaVuSans thì dùng để hiển thị tiếng Việt.
    font_name = "Helvetica"
    possible_fonts = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for font_path in possible_fonts:
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont("VietnameseFont", font_path))
            font_name = "VietnameseFont"
            break

    styles = getSampleStyleSheet()
    for style in styles.byName.values():
        style.fontName = font_name

    story = []
    story.append(Paragraph(html.escape(exam.title), styles["Title"]))
    story.append(Paragraph(f"Mức độ: {html.escape(exam.difficulty)}", styles["Normal"]))
    story.append(Paragraph("Họ và tên: ........................................", styles["Normal"]))
    story.append(Paragraph("Lớp: ...............................................", styles["Normal"]))
    story.append(Spacer(1, 12))

    for idx, q in enumerate(exam.questions, start=1):
        story.append(Paragraph(f"Câu {idx}. {html.escape(q.content)}", styles["Normal"]))
        story.append(Paragraph(f"A. {html.escape(q.option_a)}", styles["Normal"]))
        story.append(Paragraph(f"B. {html.escape(q.option_b)}", styles["Normal"]))
        story.append(Paragraph(f"C. {html.escape(q.option_c)}", styles["Normal"]))
        story.append(Paragraph(f"D. {html.escape(q.option_d)}", styles["Normal"]))
        story.append(Spacer(1, 8))

    story.append(Spacer(1, 16))
    story.append(Paragraph("Đáp án", styles["Heading2"]))
    answers = "; ".join([f"Câu {i}: {q.correct_answer}" for i, q in enumerate(exam.questions, start=1)])
    story.append(Paragraph(html.escape(answers), styles["Normal"]))

    pdf = SimpleDocTemplate(str(path), pagesize=A4)
    pdf.build(story)
    return str(path)
